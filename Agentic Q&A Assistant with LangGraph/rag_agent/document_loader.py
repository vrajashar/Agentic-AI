from pathlib import Path
from datetime import datetime
from typing import List

from langchain_core.documents import Document as LCDocument
from langchain_community.document_loaders import UnstructuredPDFLoader

from pdf2image import convert_from_path
import pytesseract

from docx import Document as DocxDocument
from llama_index.core import Document as LlamaDocument


pytesseract.pytesseract.tesseract_cmd = (
    r"C:\Program Files\Tesseract-OCR\tesseract.exe"
)

SUPPORTED_EXTENSIONS = [".pdf", ".txt", ".docx"]

# -------------------------
# PDF LOADERS (EXISTING)
# -------------------------

def load_with_unstructured(pdf_path: str) -> List[LCDocument]:
    loader = UnstructuredPDFLoader(pdf_path)
    return loader.load()


def load_with_ocr(pdf_path: str) -> List[LCDocument]:
    images = convert_from_path(pdf_path)
    documents = []

    for page_num, image in enumerate(images, start=1):
        text = pytesseract.image_to_string(image, lang="eng")

        documents.append(
            LCDocument(
                page_content=text,
                metadata={
                    "source": pdf_path,
                    "page": page_num,
                    "loader": "ocr"
                }
            )
        )

    return documents


def load_pdf(pdf_path: str) -> List[LCDocument]:
    docs = load_with_unstructured(pdf_path)

    # Fallback to OCR for scanned PDFs
    if not docs or all(len(d.page_content.strip()) < 50 for d in docs):
        docs = load_with_ocr(pdf_path)

    return docs


# -------------------------
# TXT LOADER 
# -------------------------

def load_txt(txt_path: str) -> List[LCDocument]:
    with open(txt_path, "r", encoding="utf-8", errors="ignore") as f:
        text = f.read()

    return [
        LCDocument(
            page_content=text,
            metadata={
                "source": txt_path,
                "loader": "txt"
            }
        )
    ]


# -------------------------
# DOCX LOADER 
# -------------------------

def load_docx(docx_path: str) -> List[LCDocument]:
    doc = DocxDocument(docx_path)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

    text = "\n".join(paragraphs)

    return [
        LCDocument(
            page_content=text,
            metadata={
                "source": docx_path,
                "loader": "docx"
            }
        )
    ]


# -------------------------
# CONVERSION LAYER
# -------------------------

def to_llamaindex(docs: List[LCDocument]) -> List[LlamaDocument]:
    llama_docs = []

    for doc in docs:
        llama_docs.append(
            LlamaDocument(
                text=doc.page_content,
                metadata={
                    **doc.metadata,
                    "ingested_at": datetime.utcnow().isoformat(),
                    "pipeline": "rag_agent"
                }
            )
        )

    return llama_docs


# -------------------------
# UNIFIED ENTRY POINT
# -------------------------

def load_documents(path: str) -> List[LlamaDocument]:
    path = Path(path)
    all_docs: List[LCDocument] = []

    if path.is_file():
        ext = path.suffix.lower()

        if ext == ".pdf":
            all_docs.extend(load_pdf(str(path)))
        elif ext == ".txt":
            all_docs.extend(load_txt(str(path)))
        elif ext == ".docx":
            all_docs.extend(load_docx(str(path)))
        else:
            raise ValueError(f"Unsupported file type: {ext}")

    else:
        for file in path.rglob("*"):
            if file.suffix.lower() == ".pdf":
                all_docs.extend(load_pdf(str(file)))
            elif file.suffix.lower() == ".txt":
                all_docs.extend(load_txt(str(file)))
            elif file.suffix.lower() == ".docx":
                all_docs.extend(load_docx(str(file)))

    return to_llamaindex(all_docs)
